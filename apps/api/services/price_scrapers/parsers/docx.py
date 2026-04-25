"""DOCX bulletin adapter — lazy-imports `python-docx`, feeds tables to the core.

The provincial DOCs that publish as .docx almost always embed the price
list as a single table. We read every table in the document, pick the
first one whose header row detects successfully, and hand its cells to
`table.extract_prices_from_table`.

`python-docx` is a heavy-ish dep (zipfile + lxml), so we import it
inside the function rather than at module load. Modules that never
call `parse_docx_bulletin` (e.g. unit tests for HTML-only scrapers)
don't pay the cost.
"""

from __future__ import annotations

import io
import logging
from datetime import date

from ..base import ScrapedPrice, ScrapeError
from .table import extract_effective_date, extract_prices_from_table

logger = logging.getLogger(__name__)


def parse_docx_bulletin(
    content: bytes,
    *,
    source_url: str,
    province: str,
) -> list[ScrapedPrice]:
    """Extract price rows from a DOCX bulletin's embedded tables.

    Strategy:
      * Concatenate every paragraph's text into a single string for
        date extraction (bulletins put "tháng MM/YYYY" in a title or
        decision-number paragraph *above* the table).
      * Walk every table in document order. For each, build a
        list-of-lists of cell text and delegate to the core.
      * Return the first non-empty result — subsequent tables in the
        same DOCX are usually explanatory ("notes", "references") and
        would confuse the normalizer if merged in.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover — deployed env always has it
        raise ScrapeError("python-docx not installed; cannot parse .docx bulletins") from exc

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise ScrapeError(f"could not open .docx content: {exc}") from exc

    full_text = "\n".join(p.text for p in document.paragraphs)
    effective = extract_effective_date(full_text) or date.today()

    if not document.tables:
        logger.warning("parser.docx[%s]: document has no tables", province)
        return []

    for t_idx, table in enumerate(document.tables):
        rows = _rows_from_docx_table(table)
        scraped = extract_prices_from_table(
            rows,
            effective_date=effective,
            source_url=source_url,
            province=province,
        )
        if scraped:
            logger.info(
                "parser.docx[%s]: parsed %d rows from table %d/%d",
                province,
                len(scraped),
                t_idx + 1,
                len(document.tables),
            )
            return scraped

    logger.warning(
        "parser.docx[%s]: none of %d tables parsed to any rows",
        province,
        len(document.tables),
    )
    return []


def _rows_from_docx_table(table) -> list[list[str]]:
    """Flatten a python-docx Table to list[list[str]] — cells joined by newline within each cell."""
    out: list[list[str]] = []
    for row in table.rows:
        out.append([cell.text.strip() for cell in row.cells])
    return out
