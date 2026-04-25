"""DRAWBRIDGE AI pipelines — document ingestion, RAG Q&A, conflict detection, extraction.

Pipeline overview:

  Ingestion (per upload)
    FileTypeRouter → PDFIngester | DocxIngester
      ├── text extraction (pdfplumber / python-docx)
      ├── vision analysis (GPT-4o) for drawings
      ├── table extractor (schedules)
      └── dimension parser (regex + NLP)
    → ChunkSplitter (page + semantic)
    → EmbeddingBatch (text-embedding-3-large, batch=100)
    → upsert document_chunks + embeddings
    → MetadataExtractor (title block via vision)
    → Elasticsearch index
    → emit document.ready → trigger conflict scan

  Conflict detection (cross-discipline)
    IdentifyClashCandidates → ConflictAnalyzer (LLM) → ConflictPersist → notify

  Q&A
    Query embed → hybrid retrieve (dense+BM25) → rerank → generate → bbox citations

  Extraction
    Per-document: schedules / dimensions / materials / title_block

  RFI generation
    From conflict → structured RFI draft (subject + description + related docs).
"""
from __future__ import annotations

import asyncio
import io
import logging
import os
import re
from dataclasses import dataclass
from typing import Any, Literal
from uuid import UUID, uuid4

from langchain_anthropic import ChatAnthropic
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import OpenAIEmbeddings
from langgraph.graph import END, StateGraph
from pydantic import BaseModel, Field
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

from schemas.drawbridge import (
    Conflict,
    ConflictScanResponse,
    ConflictSeverity,
    Discipline,
    ExtractedDimension,
    ExtractedMaterial,
    ExtractedSchedule,
    ExtractResponse,
    QueryResponse,
    RfiDraft,
    RfiPriority,
    ScheduleRow,
    SourceDocument,
)


# ============================================================
# Clients
# ============================================================

_ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6")
_EMBED_MODEL = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-large")
_ES_URL = os.getenv("ELASTICSEARCH_URL", "http://localhost:9200")
_RERANKER_ENDPOINT = os.getenv("RERANKER_ENDPOINT")
_VISION_MODEL = os.getenv("OPENAI_VISION_MODEL", "gpt-4o")
_RRF_K = 60
_EMBED_BATCH = 100


def _llm(temperature: float = 0.1, max_tokens: int = 4096) -> ChatAnthropic:
    return ChatAnthropic(model=_ANTHROPIC_MODEL, temperature=temperature, max_tokens=max_tokens)


def _embedder() -> OpenAIEmbeddings:
    return OpenAIEmbeddings(model=_EMBED_MODEL)


def _vec_literal(vec: list[float]) -> str:
    return "[" + ",".join(f"{x:.7f}" for x in vec) + "]"


# ============================================================
# Ingestion
# ============================================================

@dataclass
class PageBlock:
    page_number: int
    chunk_type: str  # text | table | schedule | note | dimension
    content: str
    bbox: dict[str, Any] | None = None


async def _download_bytes(storage_key: str) -> bytes:
    """Download a file from S3. Falls back to local read for dev fixtures."""
    try:
        import boto3
        from core.config import get_settings

        settings = get_settings()
        s3 = boto3.client("s3", region_name=settings.aws_region)
        obj = s3.get_object(Bucket=settings.s3_bucket, Key=storage_key)
        return obj["Body"].read()
    except Exception:
        local = os.path.join("/tmp", storage_key.replace("/", "_"))
        if os.path.exists(local):
            with open(local, "rb") as f:
                return f.read()
        return b""


def _extract_pdf_blocks(raw: bytes) -> list[PageBlock]:
    """Extract text + tables from a PDF using pdfplumber. Returns chunk candidates."""
    blocks: list[PageBlock] = []
    try:
        import pdfplumber
    except ImportError:
        return blocks

    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            txt = (page.extract_text() or "").strip()
            if txt:
                blocks.append(PageBlock(page_number=i, chunk_type="text", content=txt))
            for tbl in page.extract_tables() or []:
                rendered = "\n".join([" | ".join(str(c or "") for c in row) for row in tbl])
                if rendered.strip():
                    blocks.append(
                        PageBlock(
                            page_number=i,
                            chunk_type="schedule",
                            content=rendered,
                            bbox=None,
                        )
                    )
    return blocks


def _extract_docx_blocks(raw: bytes) -> list[PageBlock]:
    try:
        import docx  # python-docx
    except ImportError:
        return []
    d = docx.Document(io.BytesIO(raw))
    paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    if not paragraphs:
        return []
    return [PageBlock(page_number=1, chunk_type="text", content="\n".join(paragraphs))]


_DIM_REGEX = re.compile(
    r"(?P<label>[A-Za-z0-9_\- ]{0,30})\b(?P<value>\d{2,6})\s?(?P<unit>mm|cm|m)\b",
    re.IGNORECASE,
)


def _parse_dimensions(block: PageBlock) -> list[PageBlock]:
    """Pull out labeled dimensions as their own chunk_type='dimension' entries."""
    out: list[PageBlock] = []
    for m in _DIM_REGEX.finditer(block.content):
        raw = m.group(0)
        out.append(
            PageBlock(
                page_number=block.page_number,
                chunk_type="dimension",
                content=raw.strip(),
                bbox=block.bbox,
            )
        )
    return out


async def _vision_title_block(pdf_bytes: bytes) -> dict[str, Any]:
    """Use GPT-4o vision on page 1 image to read the title block. Stubbed on failure."""
    try:
        import base64

        from openai import AsyncOpenAI
        import pdfplumber
        from PIL import Image

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            if not pdf.pages:
                return {}
            img: Image.Image = pdf.pages[0].to_image(resolution=150).original  # type: ignore

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode("ascii")

        client = AsyncOpenAI()
        res = await client.chat.completions.create(
            model=_VISION_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a drafter reading an architectural/engineering drawing title block. "
                        "Extract drawing_number, title, revision, scale, discipline. "
                        "Respond with strict JSON keys: drawing_number, title, revision, scale, discipline."
                    ),
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract the title block fields."},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                    ],
                },
            ],
            response_format={"type": "json_object"},
            temperature=0.0,
        )
        import json

        return json.loads(res.choices[0].message.content or "{}")
    except Exception:
        return {}


async def _embed_batch(texts: list[str]) -> list[list[float]]:
    if not texts:
        return []
    emb = _embedder()
    # OpenAIEmbeddings.aembed_documents batches internally; we still chunk to cap token usage.
    out: list[list[float]] = []
    for i in range(0, len(texts), _EMBED_BATCH):
        chunk = texts[i : i + _EMBED_BATCH]
        out.extend(await emb.aembed_documents(chunk))
    return out


async def _es_index_chunks(
    document_id: UUID,
    project_id: UUID | None,
    organization_id: UUID,
    drawing_number: str | None,
    discipline: str | None,
    blocks: list[PageBlock],
) -> None:
    try:
        from elasticsearch import AsyncElasticsearch
    except ImportError:
        return
    es = AsyncElasticsearch(_ES_URL)
    try:
        actions = []
        for b in blocks:
            actions.append({"index": {"_index": "document_chunks"}})
            actions.append(
                {
                    "document_id": str(document_id),
                    "organization_id": str(organization_id),
                    "project_id": str(project_id) if project_id else None,
                    "drawing_number": drawing_number,
                    "discipline": discipline,
                    "page_number": b.page_number,
                    "chunk_type": b.chunk_type,
                    "content": b.content,
                }
            )
        if actions:
            await es.bulk(operations=actions, refresh=False)
    except Exception:
        pass
    finally:
        await es.close()


async def enqueue_ingest_document(
    organization_id: UUID,
    document_id: UUID,
    storage_key: str,
    mime_type: str,
) -> str | None:
    """Hand off document ingest to the arq worker.

    Returns the job id on success. If Redis is unreachable (local dev, tests),
    falls back to `asyncio.create_task` so the API still behaves correctly —
    the job just won't survive a process restart.
    """
    try:
        from workers.queue import get_pool

        pool = await get_pool()
        job = await pool.enqueue_job(
            "drawbridge_ingest_job",
            str(organization_id),
            str(document_id),
            storage_key,
            mime_type,
        )
        return job.job_id if job else None
    except Exception:  # pragma: no cover — graceful degradation for local dev
        logger.warning(
            "drawbridge ingest: arq unavailable, running inline task for doc=%s",
            document_id,
        )
        asyncio.create_task(
            _ingest_document(
                organization_id=organization_id,
                document_id=document_id,
                storage_key=storage_key,
                mime_type=mime_type,
            )
        )
        return None


async def _ingest_document(
    organization_id: UUID,
    document_id: UUID,
    storage_key: str,
    mime_type: str,
) -> None:
    """End-to-end ingest. Persists chunks + embeddings, updates document.processing_status."""
    from db.session import TenantAwareSession
    from models.drawbridge import Document as DocumentModel

    async with TenantAwareSession(organization_id) as session:
        doc = await session.get(DocumentModel, document_id)
        if doc is None:
            return
        doc.processing_status = "processing"
        await session.commit()

        raw = await _download_bytes(storage_key)

        blocks: list[PageBlock]
        if "pdf" in (mime_type or "") or storage_key.lower().endswith(".pdf"):
            blocks = _extract_pdf_blocks(raw)
            # Augment text blocks with dimension extractions.
            aug: list[PageBlock] = []
            for b in blocks:
                if b.chunk_type == "text":
                    aug.extend(_parse_dimensions(b))
            blocks.extend(aug)
            if not doc.drawing_number or not doc.title:
                tb = await _vision_title_block(raw)
                if tb:
                    doc.drawing_number = doc.drawing_number or tb.get("drawing_number")
                    doc.title = doc.title or tb.get("title")
                    doc.revision = doc.revision or tb.get("revision")
                    doc.scale = doc.scale or tb.get("scale")
                    doc.discipline = doc.discipline or tb.get("discipline")
        elif "word" in (mime_type or "") or storage_key.lower().endswith(".docx"):
            blocks = _extract_docx_blocks(raw)
        else:
            blocks = []

        if not blocks:
            doc.processing_status = "ready"
            await session.commit()
            return

        # Embed + insert chunks.
        embeddings = await _embed_batch([b.content for b in blocks])
        for b, vec in zip(blocks, embeddings):
            await session.execute(
                text(
                    """
                    INSERT INTO document_chunks
                      (id, document_id, organization_id, project_id,
                       page_number, chunk_type, content, bbox, embedding)
                    VALUES
                      (:id, :document_id, :organization_id, :project_id,
                       :page_number, :chunk_type, :content,
                       CAST(:bbox AS JSONB), CAST(:embedding AS vector))
                    """
                ),
                {
                    "id": str(uuid4()),
                    "document_id": str(document_id),
                    "organization_id": str(organization_id),
                    "project_id": str(doc.project_id) if doc.project_id else None,
                    "page_number": b.page_number,
                    "chunk_type": b.chunk_type,
                    "content": b.content,
                    "bbox": None if b.bbox is None else __import__("json").dumps(b.bbox),
                    "embedding": _vec_literal(vec),
                },
            )

        doc.processing_status = "ready"
        await session.commit()

        # Index in ES + trigger conflict scan for fresh drawings.
        await _es_index_chunks(
            document_id=document_id,
            project_id=doc.project_id,
            organization_id=organization_id,
            drawing_number=doc.drawing_number,
            discipline=doc.discipline,
            blocks=blocks,
        )

        if doc.doc_type == "drawing" and doc.project_id:
            try:
                await run_conflict_scan(
                    db=session,
                    organization_id=organization_id,
                    project_id=doc.project_id,
                    document_ids=[document_id],
                    severities=None,
                    raised_by=None,
                )
            except Exception:
                pass


# ============================================================
# Q&A
# ============================================================

_QA_SYSTEM = """You are DRAWBRIDGE, an assistant that answers questions about AEC project documents.

Rules:
- Answer strictly from the provided context chunks.
- Every factual claim MUST cite one or more chunks by their [index] and drawing number.
- If the context is insufficient, say so and lower confidence.
- Respond in the user's language: {language}.

Output MUST be valid JSON with:
{{
  "answer": string,
  "confidence": number in [0,1],
  "citations": [{{
    "chunk_index": int,
    "drawing_number": string,
    "page": int,
    "excerpt": string
  }}],
  "related_questions": [string, string, string]
}}
"""

_QA_USER = """Question: {question}

Context chunks:
{context}

Return JSON only."""


def _detect_language(s: str) -> str:
    vi_markers = "ăâđêôơưạảấầẩẫậắằẳẵặẹẻẽếềểễệỉịọỏốồổỗộớờởỡợụủứừửữựỳỵỷỹ"
    return "vi" if any(c in s.lower() for c in vi_markers) else "en"


async def _dense_search_chunks(
    db: AsyncSession,
    project_id: UUID,
    vec_query: list[float],
    disciplines: list[Discipline] | None,
    document_ids: list[UUID] | None,
    top_k: int,
) -> list[dict[str, Any]]:
    params: dict[str, Any] = {
        "vec": _vec_literal(vec_query),
        "project_id": str(project_id),
        "limit": top_k * 3,
    }
    where = ["c.project_id = :project_id"]
    if disciplines:
        where.append("d.discipline = ANY(:disciplines)")
        params["disciplines"] = [x.value for x in disciplines]
    if document_ids:
        where.append("c.document_id = ANY(:document_ids)")
        params["document_ids"] = [str(i) for i in document_ids]

    # Use `embedding_half halfvec(3072)` + HNSW index (added in 0007_drawbridge_hnsw).
    # `vector(3072)` has no supported ANN index (pgvector caps at 2000 dims).
    sql = text(
        f"""
        SELECT c.id AS chunk_id, c.document_id, c.page_number, c.chunk_type,
               c.content, c.bbox, d.drawing_number, d.title, d.discipline,
               1 - (c.embedding_half <=> CAST(:vec AS halfvec)) AS score
        FROM document_chunks c
        JOIN documents d ON d.id = c.document_id
        WHERE {" AND ".join(where)}
        ORDER BY c.embedding_half <=> CAST(:vec AS halfvec)
        LIMIT :limit
        """
    )
    rows = (await db.execute(sql, params)).mappings().all()
    return [dict(r) for r in rows]


async def _sparse_search_chunks(
    project_id: UUID,
    question: str,
    disciplines: list[Discipline] | None,
    top_k: int,
) -> list[dict[str, Any]]:
    try:
        from elasticsearch import AsyncElasticsearch
    except ImportError:
        return []
    es = AsyncElasticsearch(_ES_URL)
    try:
        must: list[dict[str, Any]] = [{"match": {"content": question}}]
        filters: list[dict[str, Any]] = [{"term": {"project_id": str(project_id)}}]
        if disciplines:
            filters.append({"terms": {"discipline": [d.value for d in disciplines]}})
        res = await es.search(
            index="document_chunks",
            size=top_k * 3,
            query={"bool": {"must": must, "filter": filters}},
        )
        return [
            {
                "chunk_id": h["_id"],
                "document_id": h["_source"].get("document_id"),
                "page_number": h["_source"].get("page_number"),
                "chunk_type": h["_source"].get("chunk_type"),
                "content": h["_source"].get("content"),
                "drawing_number": h["_source"].get("drawing_number"),
                "discipline": h["_source"].get("discipline"),
                "score": h["_score"],
                "bbox": None,
                "title": None,
            }
            for h in res.get("hits", {}).get("hits", [])
        ]
    except Exception:
        return []
    finally:
        await es.close()


def _rrf(dense: list[dict[str, Any]], sparse: list[dict[str, Any]], k: int = _RRF_K) -> list[dict[str, Any]]:
    scores: dict[str, float] = {}
    payloads: dict[str, dict[str, Any]] = {}
    for rank, it in enumerate(dense):
        key = str(it["chunk_id"])
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank + 1)
        payloads[key] = it
    for rank, it in enumerate(sparse):
        key = str(it["chunk_id"])
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank + 1)
        payloads.setdefault(key, it)
    return sorted(payloads.values(), key=lambda p: scores[str(p["chunk_id"])], reverse=True)


async def _rerank(question: str, candidates: list[dict[str, Any]], top_k: int) -> list[dict[str, Any]]:
    if not _RERANKER_ENDPOINT or not candidates:
        return candidates[:top_k]
    try:
        import httpx

        async with httpx.AsyncClient(timeout=10.0) as client:
            res = await client.post(
                _RERANKER_ENDPOINT,
                json={"query": question, "documents": [c["content"] or "" for c in candidates]},
            )
            res.raise_for_status()
            order = res.json()["ranked_indices"]
            return [candidates[i] for i in order[:top_k]]
    except Exception:
        return candidates[:top_k]


class _QaState(BaseModel):
    question: str
    language: str
    project_id: UUID
    disciplines: list[Discipline] | None = None
    document_ids: list[UUID] | None = None
    top_k: int = 12
    candidates: list[dict[str, Any]] = Field(default_factory=list)
    answer: QueryResponse | None = None


async def answer_document_query(
    db: AsyncSession,
    organization_id: UUID,
    project_id: UUID,
    question: str,
    disciplines: list[Discipline] | None,
    document_ids: list[UUID] | None,
    top_k: int,
    language: str | None,
) -> QueryResponse:
    lang = language or _detect_language(question)

    graph = StateGraph(_QaState)

    async def retrieve(state: _QaState) -> _QaState:
        vec = await _embedder().aembed_query(state.question)
        dense = await _dense_search_chunks(
            db, state.project_id, vec, state.disciplines, state.document_ids, state.top_k
        )
        sparse = await _sparse_search_chunks(state.project_id, state.question, state.disciplines, state.top_k)
        fused = _rrf(dense, sparse)
        state.candidates = await _rerank(state.question, fused, state.top_k)
        return state

    async def generate(state: _QaState) -> _QaState:
        context_lines = [
            f"[{i}] {c.get('drawing_number') or '?'} (p{c.get('page_number') or '?'}, "
            f"{c.get('discipline') or 'unknown'})\n{c.get('content') or ''}"
            for i, c in enumerate(state.candidates)
        ]
        context = "\n\n".join(context_lines) or "(no documents found)"
        prompt = ChatPromptTemplate.from_messages(
            [
                ("system", _QA_SYSTEM.format(language=state.language)),
                ("human", _QA_USER),
            ]
        )
        parser = JsonOutputParser()
        chain = prompt | _llm(temperature=0.1) | parser
        raw = await chain.ainvoke({"question": state.question, "context": context})

        sources: list[SourceDocument] = []
        for cit in raw.get("citations", []) or []:
            idx = cit.get("chunk_index")
            if idx is None or not (0 <= int(idx) < len(state.candidates)):
                continue
            c = state.candidates[int(idx)]
            sources.append(
                SourceDocument(
                    document_id=UUID(str(c["document_id"])) if c.get("document_id") else uuid4(),
                    drawing_number=c.get("drawing_number") or cit.get("drawing_number"),
                    title=c.get("title"),
                    discipline=Discipline(c["discipline"]) if c.get("discipline") else None,
                    page=c.get("page_number") or cit.get("page"),
                    excerpt=cit.get("excerpt") or (c.get("content") or "")[:400],
                    bbox=c.get("bbox"),
                )
            )
        state.answer = QueryResponse(
            answer=str(raw.get("answer", "")),
            confidence=float(raw.get("confidence", 0.5) or 0.5),
            source_documents=sources,
            related_questions=[str(q) for q in (raw.get("related_questions") or [])][:3],
        )
        return state

    graph.add_node("retrieve", retrieve)
    graph.add_node("generate", generate)
    graph.set_entry_point("retrieve")
    graph.add_edge("retrieve", "generate")
    graph.add_edge("generate", END)
    app = graph.compile()

    final: _QaState = await app.ainvoke(
        _QaState(
            question=question,
            language=lang,
            project_id=project_id,
            disciplines=disciplines,
            document_ids=document_ids,
            top_k=top_k,
        )
    )
    assert final.answer is not None
    return final.answer


# ============================================================
# Conflict detection
# ============================================================

_CROSS_DISCIPLINE_PAIRS = {
    ("architectural", "structural"),
    ("structural", "mep"),
    ("architectural", "mep"),
    ("structural", "civil"),
    ("architectural", "civil"),
}

_CONFLICT_SYSTEM = """You are DRAWBRIDGE, reviewing two drawing extracts from different disciplines.

Task: identify any SPECIFIC conflict in dimensions, materials, elevations, structural requirements,
or clearances. Only flag real conflicts — do not invent. If the extracts simply describe different
parts of the building without overlap, return is_conflict=false.

Return strict JSON:
{
  "is_conflict": boolean,
  "severity": "critical" | "major" | "minor" | null,
  "type": "dimension" | "material" | "structural" | "elevation" | null,
  "explanation": string,
  "short_description": string
}
"""


async def run_conflict_scan(
    db: AsyncSession,
    organization_id: UUID,
    project_id: UUID,
    document_ids: list[UUID] | None,
    severities: list[ConflictSeverity] | None,
    raised_by: UUID | None,
) -> ConflictScanResponse:
    """Scan new-or-specified documents against the rest of the project for cross-discipline conflicts."""
    from models.drawbridge import Conflict as ConflictModel

    # Pick candidate-pool chunks with a dimension / schedule / note character,
    # filtered to target documents first then the rest of the project.
    target_where = "c.organization_id = :org AND c.project_id = :project"
    params: dict[str, Any] = {"org": str(organization_id), "project": str(project_id)}
    if document_ids:
        target_where += " AND c.document_id = ANY(:doc_ids)"
        params["doc_ids"] = [str(i) for i in document_ids]

    target_sql = text(
        f"""
        SELECT c.id, c.document_id, c.content, c.page_number, c.chunk_type,
               c.embedding::text AS emb, d.discipline, d.drawing_number
        FROM document_chunks c
        JOIN documents d ON d.id = c.document_id
        WHERE {target_where}
          AND c.chunk_type IN ('dimension', 'schedule', 'note', 'text')
          AND c.embedding IS NOT NULL
        """
    )
    targets = (await db.execute(target_sql, params)).mappings().all()

    scanned_documents = len({t["document_id"] for t in targets})
    candidates_evaluated = 0
    persisted: list[Conflict] = []

    for t in targets:
        # Retrieve nearest neighbors from OTHER disciplines in the same project.
        neighbor_sql = text(
            """
            SELECT c.id, c.document_id, c.content, c.page_number, c.chunk_type,
                   d.discipline, d.drawing_number,
                   1 - (c.embedding_half <=> CAST(:emb AS halfvec)) AS score
            FROM document_chunks c
            JOIN documents d ON d.id = c.document_id
            WHERE c.organization_id = :org
              AND c.project_id = :project
              AND d.discipline IS NOT NULL
              AND d.discipline <> :own_discipline
              AND c.document_id <> :own_doc
              AND c.embedding IS NOT NULL
            ORDER BY c.embedding_half <=> CAST(:emb AS halfvec)
            LIMIT 5
            """
        )
        neighbors = (
            await db.execute(
                neighbor_sql,
                {
                    "emb": t["emb"],
                    "org": str(organization_id),
                    "project": str(project_id),
                    "own_discipline": t["discipline"] or "",
                    "own_doc": str(t["document_id"]),
                },
            )
        ).mappings().all()

        for n in neighbors:
            pair = tuple(sorted([t["discipline"] or "", n["discipline"] or ""]))
            if pair not in _CROSS_DISCIPLINE_PAIRS:
                continue
            candidates_evaluated += 1

            verdict = await _analyze_conflict_pair(
                a=dict(t),
                b=dict(n),
            )
            if not verdict.get("is_conflict"):
                continue

            severity = verdict.get("severity") or "minor"
            if severities and severity not in [s.value for s in severities]:
                continue

            row = ConflictModel(
                id=uuid4(),
                organization_id=organization_id,
                project_id=project_id,
                status="open",
                severity=severity,
                conflict_type=verdict.get("type"),
                description=verdict.get("short_description"),
                document_a_id=t["document_id"],
                chunk_a_id=t["id"],
                document_b_id=n["document_id"],
                chunk_b_id=n["id"],
                ai_explanation=verdict.get("explanation"),
            )
            db.add(row)
            await db.flush()
            persisted.append(Conflict.model_validate(row))

    await db.commit()

    # Notify (stub — wire to email/in-app in your notification service).
    if persisted and raised_by:
        await _notify_conflict_created(organization_id, project_id, len(persisted))

    return ConflictScanResponse(
        project_id=project_id,
        scanned_documents=scanned_documents,
        candidates_evaluated=candidates_evaluated,
        conflicts_found=len(persisted),
        conflicts=persisted,
    )


async def _analyze_conflict_pair(a: dict[str, Any], b: dict[str, Any]) -> dict[str, Any]:
    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", _CONFLICT_SYSTEM),
            (
                "human",
                "Extract A (discipline={da}, drawing={na}, page={pa}):\n{ca}\n\n"
                "Extract B (discipline={db_}, drawing={nb}, page={pb}):\n{cb}\n\nReturn JSON only.",
            ),
        ]
    )
    chain = prompt | _llm(temperature=0.0) | JsonOutputParser()
    try:
        return await chain.ainvoke(
            {
                "da": a.get("discipline"),
                "na": a.get("drawing_number"),
                "pa": a.get("page_number"),
                "ca": (a.get("content") or "")[:2000],
                "db_": b.get("discipline"),
                "nb": b.get("drawing_number"),
                "pb": b.get("page_number"),
                "cb": (b.get("content") or "")[:2000],
            }
        )
    except Exception:
        return {"is_conflict": False}


async def _notify_conflict_created(org_id: UUID, project_id: UUID, count: int) -> None:
    """Placeholder for email + in-app notification."""
    return None


# ============================================================
# Extraction
# ============================================================

_EXTRACT_SYSTEM = """Extract structured data from this drawing/spec text.

Return strict JSON:
{
  "schedules": [{"name": str, "page": int|null, "columns": [str], "rows": [{"cells": {col: value}}]}],
  "dimensions": [{"label": str, "value_mm": number|null, "raw": str, "page": int|null}],
  "materials":  [{"code": str|null, "description": str, "quantity": number|null, "unit": str|null, "page": int|null}],
  "title_block": {"drawing_number": str|null, "title": str|null, "revision": str|null, "scale": str|null, "discipline": str|null}
}
"""


async def extract_document_data(
    db: AsyncSession,
    document_id: UUID,
    target: Literal["schedule", "dimensions", "materials", "title_block", "all"],
    pages: list[int] | None,
) -> ExtractResponse:
    from models.drawbridge import Document as DocumentModel

    sql = text(
        """
        SELECT c.page_number, c.chunk_type, c.content
        FROM document_chunks c
        WHERE c.document_id = :doc_id
        ORDER BY c.page_number
        """
    )
    rows = (await db.execute(sql, {"doc_id": str(document_id)})).mappings().all()
    if pages:
        rows = [r for r in rows if r["page_number"] in pages]

    doc = await db.get(DocumentModel, document_id)
    header = f"Document {doc.drawing_number or document_id} ({doc.discipline or '?'})" if doc else ""

    body = "\n\n".join(
        f"[p{r['page_number']}, {r['chunk_type']}] {r['content']}"
        for r in rows
        if r["content"]
    ) or "(no content)"

    prompt = ChatPromptTemplate.from_messages(
        [("system", _EXTRACT_SYSTEM), ("human", f"{header}\nTarget: {target}\n\n{body}")]
    )
    parser = JsonOutputParser()
    try:
        parsed = await (prompt | _llm(temperature=0.0, max_tokens=6000) | parser).ainvoke({})
    except Exception:
        parsed = {}

    schedules = [
        ExtractedSchedule(
            name=str(s.get("name", "Schedule")),
            page=s.get("page"),
            columns=[str(c) for c in s.get("columns", [])],
            rows=[ScheduleRow(cells=r.get("cells", {})) for r in s.get("rows", [])],
        )
        for s in parsed.get("schedules", []) or []
    ]
    dimensions = [
        ExtractedDimension(
            label=str(d.get("label", "")),
            value_mm=d.get("value_mm"),
            raw=str(d.get("raw", "")),
            page=d.get("page"),
            bbox=d.get("bbox"),
        )
        for d in parsed.get("dimensions", []) or []
    ]
    materials = [
        ExtractedMaterial(
            code=m.get("code"),
            description=str(m.get("description", "")),
            quantity=m.get("quantity"),
            unit=m.get("unit"),
            page=m.get("page"),
        )
        for m in parsed.get("materials", []) or []
    ]
    return ExtractResponse(
        document_id=document_id,
        schedules=schedules,
        dimensions=dimensions,
        materials=materials,
        title_block=parsed.get("title_block"),
    )


# ============================================================
# RFI draft generation
# ============================================================

_RFI_SYSTEM = """You are an experienced project engineer drafting an RFI (Request for Information).

Given a construction drawing conflict, produce a professional RFI:
- subject: <= 80 chars, specific
- description: reference both drawings explicitly, state the conflict, ask a clear question
- include unit conversions when helpful
- do NOT speculate on the answer

Return strict JSON: {"subject": str, "description": str}
"""


async def draft_rfi_from_conflict(db: AsyncSession, conflict_id: UUID) -> RfiDraft:
    from models.drawbridge import (
        Conflict as ConflictModel,
        Document as DocumentModel,
        DocumentChunk as DocumentChunkModel,
    )

    conflict = await db.get(ConflictModel, conflict_id)
    if conflict is None:
        raise ValueError("Conflict not found")

    doc_a = await db.get(DocumentModel, conflict.document_a_id) if conflict.document_a_id else None
    doc_b = await db.get(DocumentModel, conflict.document_b_id) if conflict.document_b_id else None
    chunk_a = await db.get(DocumentChunkModel, conflict.chunk_a_id) if conflict.chunk_a_id else None
    chunk_b = await db.get(DocumentChunkModel, conflict.chunk_b_id) if conflict.chunk_b_id else None

    user = (
        f"Conflict type: {conflict.conflict_type or 'unspecified'}\n"
        f"Severity: {conflict.severity or 'unspecified'}\n"
        f"Description: {conflict.description or ''}\n"
        f"AI explanation: {conflict.ai_explanation or ''}\n\n"
        f"Drawing A ({doc_a.drawing_number if doc_a else '?'}, {doc_a.discipline if doc_a else '?'}):\n"
        f"{(chunk_a.content if chunk_a else '') or ''}\n\n"
        f"Drawing B ({doc_b.drawing_number if doc_b else '?'}, {doc_b.discipline if doc_b else '?'}):\n"
        f"{(chunk_b.content if chunk_b else '') or ''}"
    )

    prompt = ChatPromptTemplate.from_messages([("system", _RFI_SYSTEM), ("human", user)])
    parser = JsonOutputParser()
    try:
        raw = await (prompt | _llm(temperature=0.1) | parser).ainvoke({})
    except Exception:
        raw = {
            "subject": f"Clarification needed: {conflict.conflict_type or 'conflict'} between drawings",
            "description": conflict.ai_explanation or conflict.description or "Please review and advise.",
        }

    related: list[UUID] = []
    if conflict.document_a_id:
        related.append(conflict.document_a_id)
    if conflict.document_b_id:
        related.append(conflict.document_b_id)

    return RfiDraft(
        subject=str(raw.get("subject", "RFI")).strip()[:200],
        description=str(raw.get("description", "")).strip(),
        related_document_ids=related,
        priority=RfiPriority.high,
    )
